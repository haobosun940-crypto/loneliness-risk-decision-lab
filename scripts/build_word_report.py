#!/usr/bin/env python3
"""Build the Word research report for the loneliness risk project."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
ASSET_DIR = ROOT / "assets"
OUTPUT_DIR = ROOT / "outputs"
DOCX_PATH = OUTPUT_DIR / "loneliness_risk_research_report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(18, 32, 51)
MUTED = RGBColor(85, 94, 108)
TEAL = RGBColor(15, 118, 110)
RED = RGBColor(178, 58, 72)
GOLD = RGBColor(170, 103, 0)
LIGHT_FILL = "F4F6F9"
PALE_BLUE = "E8EEF5"


def load_json(name: str):
    return json.loads((DATA_DIR / name).read_text(encoding="utf-8"))


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, inches: float) -> None:
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def format_cell(cell, header=False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if header:
        shade_cell(cell, PALE_BLUE)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.12
        for run in paragraph.runs:
            set_run_font(run, size=9.5, color=INK, bold=header)


def add_table(doc, rows, widths, header_fill=PALE_BLUE):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = str(value)
            set_cell_width(cell, widths[c_idx])
            if r_idx == 0:
                shade_cell(cell, header_fill)
            format_cell(cell, header=r_idx == 0)
    doc.add_paragraph()
    return table


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)


def add_body(doc, text: str):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Body Text"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_bullet(doc, text: str):
    paragraph = doc.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    return paragraph


def add_number(doc, text: str):
    paragraph = doc.add_paragraph(text, style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    return paragraph


def add_figure(doc, filename: str, caption: str, width=6.15):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(ASSET_DIR / filename), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.style = doc.styles["Caption"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cap


def set_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    body = doc.styles["Body Text"]
    body.font.name = "Calibri"
    body._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    body._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    body.font.size = Pt(11)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(8)
    body.paragraph_format.line_spacing = 1.333

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.color.rgb = BLUE
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.color.rgb = DARK_BLUE
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = "Loneliness & Risk Decision Lab"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Page ")
    add_page_number(footer)
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)


def build_docx() -> None:
    stats = load_json("stats_summary.json")
    sources = load_json("sources.json")
    config = load_json("study_config.json")
    links = load_json("publication_links.json")

    OUTPUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    set_styles(doc)

    kicker = doc.add_paragraph("QUANTITATIVE PSYCHOLOGY AND STATISTICS REPORT")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.runs[0]
    set_run_font(run, size=10, color=TEAL, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Loneliness, Social Connection, and Risk Decisions")
    set_run_font(title_run, size=28, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "A research-product system for measuring how perceived isolation relates to immediate rewards, impulsive spending, risky choice, and conflict response."
    )
    set_run_font(subtitle_run, size=12, color=MUTED)

    developer = doc.add_paragraph()
    developer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    developer_run = developer.add_run("Developed by He Haoze (何昊泽)")
    set_run_font(developer_run, size=11, color=TEAL, bold=True)

    meta_rows = [
        ["Field", "Value"],
        ["Developer", "He Haoze (何昊泽)"],
        ["Target population", "Late adolescents and emerging adults, ages 16-24"],
        ["Data status", f"{stats['n']} synthetic pilot rows plus separated live submissions"],
        ["Current public survey", links["current_public_survey"]],
        ["GitHub repository", links["github_repository"]],
        ["Permanent deployment path", links["render_blueprint"]],
    ]
    add_table(doc, meta_rows, [1.75, 4.75], header_fill=LIGHT_FILL)

    doc.add_heading("Abstract", level=1)
    add_body(
        doc,
        "This report presents a complete research design and pilot analysis for the question of whether loneliness predicts risk-oriented decision making. The study boundary combines social psychology, behavioral decision science, and applied statistics. The predictor is perceived loneliness, interpreted alongside social connectedness as a possible buffer. The outcome is a composite Risk Decision Index built from delay-choice tasks, impulsive-spending items, high-risk choices, and conflict-response scenarios. The system includes a public website, survey route, scoring API, SQLite database, statistical dashboard, PDF paper, Word report, Excel workbook, slide deck, video asset, and reproducible source repository.",
    )

    doc.add_heading("Research Boundary", level=1)
    add_body(
        doc,
        "The research follows the boundary formula of broad field plus specific population plus specific variable plus specific outcome plus specific context. This keeps the topic narrow enough for quantitative analysis while still preserving a clear psychological mechanism.",
    )
    add_table(
        doc,
        [
            ["Boundary element", "Operational definition"],
            ["Broad field", "Social psychology, behavioral decision science, and applied statistics"],
            ["Specific population", "Late adolescents and emerging adults in school, college, and early-career peer networks"],
            ["Specific variable", "Loneliness score with social connectedness as a theoretically important buffer"],
            ["Specific outcome", "Risk Decision Index, scaled from 0 to 100"],
            ["Specific context", "Digitally mediated peer communication and real-world social contact"],
        ],
        [1.65, 4.85],
    )

    doc.add_heading("Theory and Hypotheses", level=1)
    add_body(
        doc,
        "Loneliness is treated as perceived social isolation rather than simple physical solitude. The psychological mechanism is that perceived isolation may increase social threat monitoring, urgency, rejection sensitivity, and short-horizon regulation. Those mechanisms can plausibly shift reward timing, spending, risk tolerance, and conflict interpretation.",
    )
    for hypothesis in config["questionnaire"]["hypotheses"]:
        add_bullet(doc, hypothesis)

    doc.add_heading("Measurement and Experiment Design", level=1)
    add_body(
        doc,
        "The questionnaire combines self-report indicators and short behavioral choice tasks. The self-report measures estimate loneliness, connection, digital contact, offline contact, stress, and sleep. The tasks estimate immediate-reward bias, impulsive spending, high-risk choice, and defensive conflict response. Each construct is converted into a 0-100 score so the dashboard, workbook, and model outputs can be compared on a common scale.",
    )
    construct_rows = [["Construct", "Measurement role"]]
    for item in config["questionnaire"]["constructs"][:8]:
        label = item.split("(")[0].strip()
        construct_rows.append([label, item])
    add_table(doc, construct_rows, [1.7, 4.8])
    add_figure(doc, "method_flow.png", "Figure 1. Research system workflow from question boundary to questionnaire, scoring, database, models, and reporting.", width=6.1)

    doc.add_heading("Statistical Analysis Plan", level=1)
    add_body(
        doc,
        "The analysis has four layers: reliability, correlation, ANOVA, and regression. Reliability checks whether the custom loneliness and social-connection batteries are internally coherent. Correlations describe zero-order relationships. ANOVA tests whether low, moderate, and high loneliness groups differ in Risk Decision Index. OLS regression estimates whether loneliness remains associated with risk after social connection, stress, age, sleep, platform sessions, and the loneliness-by-connection interaction are included.",
    )
    add_table(
        doc,
        [
            ["Analysis", "Purpose"],
            ["Cronbach alpha", "Check internal consistency of psychological item batteries"],
            ["Correlation matrix", "Show pairwise associations among constructs"],
            ["ANOVA", "Compare Risk Decision Index across loneliness tertiles"],
            ["OLS regression", "Estimate the controlled association between loneliness and risk"],
            ["Logistic tendency model", "Model top-quartile risk classification as an odds outcome"],
        ],
        [1.65, 4.85],
    )

    doc.add_heading("Results", level=1)
    alpha = stats["cronbach_alpha"]
    add_body(
        doc,
        f"The synthetic pilot shows strong internal consistency: alpha={alpha['loneliness_custom_7_item']} for the loneliness battery and alpha={alpha['social_connection_custom_6_item']} for the social-connection battery. ANOVA yields F({stats['anova']['df_between']}, {stats['anova']['df_within']})={stats['anova']['F']:.2f}, eta squared={stats['anova']['eta_squared']:.3f}, and permutation p={stats['anova']['p_permutation']:.4f}. In this demonstration dataset, the group difference is large and directionally consistent with the hypothesis.",
    )
    add_figure(doc, "chart_anova_groups.png", "Figure 2. Risk Decision Index rises across low, moderate, and high loneliness groups.", width=5.9)
    add_figure(doc, "chart_regression_coefficients.png", "Figure 3. OLS coefficients separate loneliness from social connection, stress, and controls.", width=5.9)

    ols_rows = [["Term", "Estimate", "SE", "p approx"]]
    for row in stats["ols"]:
        ols_rows.append([
            row["term"].replace("_", " "),
            f"{row['estimate']:.2f}",
            f"{row['std_error']:.2f}",
            "<.001" if row["p_approx"] < 0.001 else f"{row['p_approx']:.3f}",
        ])
    add_table(doc, ols_rows, [2.75, 1.15, 1.15, 1.45])

    doc.add_heading("Interpreting the Visualizations", level=1)
    add_body(
        doc,
        "The charts should be read as evidence tools, not decoration. The ANOVA bar chart communicates group separation. The regression coefficient chart communicates controlled association. The scatterplot communicates individual-level variation around the trend, which is important because psychological prediction is probabilistic rather than deterministic.",
    )
    add_figure(doc, "chart_scatter_loneliness_risk.png", "Figure 4. Individual-level scatterplot showing the loneliness-risk pattern and residual variation.", width=5.9)
    add_figure(doc, "chart_decision_types.png", "Figure 5. Decision-type labels summarize behavioral tendency and are not clinical categories.", width=5.3)

    doc.add_heading("Public Website and Data Pipeline", level=1)
    add_body(
        doc,
        "The website is designed as the live research instrument. Participants enter the survey at /survey, receive a personal risk-decision profile, and can submit the response into SQLite. The dashboard route shows summary statistics, recent live submissions, CSV export, and model outputs. The data pipeline keeps synthetic pilot rows and live submissions separated so demonstration data cannot be mistaken for actual human-subject observations.",
    )
    add_table(
        doc,
        [
            ["Route or artifact", "Function"],
            ["/survey", "Public questionnaire and personal report generation"],
            ["/dashboard", "Live data preview, QR sharing, CSV export, and statistical cockpit"],
            ["/research", "Workflow, model outputs, ANOVA, and interpretation"],
            ["/downloads", "PDF, Word, PPTX, workbook, video, CSV export, and full package"],
            ["GitHub", links["github_repository"]],
            ["Render Blueprint", links["render_blueprint"]],
        ],
        [1.75, 4.75],
    )

    doc.add_heading("Ethics, Data Disclosure, and Limitations", level=1)
    for item in [
        "The seeded pilot rows are synthetic demonstration data, not collected human-subject observations.",
        "Live submissions should be used only with clear consent language and minimal identifying information.",
        "The Risk Decision Index is educational and statistical; it is not a clinical diagnosis.",
        "The model is cross-sectional and cannot prove that loneliness causes risk decisions.",
        "A stronger next study should collect a real sample, predefine inclusion criteria, and rerun reliability, ANOVA, and regression on live data only.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Conclusion", level=1)
    add_body(
        doc,
        "This project demonstrates how a broad psychological topic can become a rigorous quantitative research system. The core contribution is the integration of theory, measurement, database design, live survey collection, statistical modeling, and publishable reporting. In the synthetic pilot, loneliness is positively associated with Risk Decision Index, while social connectedness is negatively associated with risk. The next scientific step is to collect real responses through the public survey, preserve transparent data-source labels, and update the report when live data are sufficient for empirical claims.",
    )

    doc.add_heading("References and Public Links", level=1)
    link_paragraphs = [
        ("Current public survey: ", links["current_public_survey"]),
        ("GitHub repository: ", links["github_repository"]),
        ("Render Blueprint: ", links["render_blueprint"]),
    ]
    for label, url in link_paragraphs:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(label)
        set_run_font(run, size=10.5, color=INK, bold=True)
        add_hyperlink(paragraph, url, url)
    doc.add_paragraph()
    for idx, src in enumerate(sources, 1):
        paragraph = doc.add_paragraph(style="Body Text")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"{idx}. {src['citation']} URL: {src['url']}")
        set_run_font(run, size=9.5, color=INK)

    doc.core_properties.title = "Loneliness, Social Connection, and Risk Decisions"
    doc.core_properties.author = "He Haoze (何昊泽)"
    doc.core_properties.subject = "Quantitative psychology and statistics research report"
    doc.core_properties.keywords = "loneliness, social connectedness, risk decisions, ANOVA, regression, survey"
    doc.core_properties.comments = "Developed by He Haoze (何昊泽). Synthetic and live data are explicitly separated."
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
